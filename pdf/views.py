from distutils import errors
import json
from os import remove
import os
import re
import uuid
from headless_browser.headless_browser import takeScreenShot

from rest_framework import viewsets
from rest_framework.views import APIView
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from reverseGeoCodeing.settings import BASE_DIR, STATIC_ROOT
from utils.line_converter import convertLatLngLine, convertLatLngLineFake
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from docx.shared import Cm

import polyline

class ConvertPdfView(APIView):

    def post(self,request, format=None):
        withPicturs=json.loads(request.data["withPictures"])
        
        from pdf2docx import Converter,parse
      
        file=request.FILES['file']
        randomName = str(uuid.uuid4())
        pdffilename="static/"+randomName+".pdf"
        pdf_file = default_storage.save(pdffilename, ContentFile(file.read()))

        randomName = str(uuid.uuid4())
        docx_file = "static/"+randomName+".docx"  
        # convert pdf to docx
        cv = Converter(pdf_file)
        cv.convert(docx_file, multi_processing=True)      # all pages by default
        cv.close()
        # parse(pdf_file,docx_file, multi_processing=True)
        os.remove(pdf_file)
        from docx import Document
        from docx.enum.table import WD_ROW_HEIGHT_RULE

        doc = Document(docx_file)
 



     



        
            

        
        path=[]
        for k,table in enumerate(doc.tables):
            for i,row in enumerate(table.rows):
                for j,cell in enumerate(row.cells):
                        p = cell.paragraphs[0]
                        if(True or(i==2 and (j==1 or j==3) and k==5)):
                            # print (p._p.xml)

                           
                            
                            latlongs=GetParagraphText(p)
                            try:
                                if latlongs is not None and withPicturs: path.append(latlongs)
                            except:
                                pass
                            

                        if withPicturs and re.search("Au total",p.text) and j==0:
                            # do something with path 
                            encodedPolyline=polyline.encode(path, 5)
                            path=[]
                            row.height =Cm(5)
                            paragraph=cell.add_paragraph()

                            print("waaah" ,i," ",j) 
                            image=takeScreenShot(encodedPolyline)
                            run = paragraph.add_run()
                            run.add_picture(image,height =Cm(7))
                row.height_rule = WD_ROW_HEIGHT_RULE.AUTO            
         
        

        os.remove(docx_file)
        
        randomName = str(uuid.uuid4())
        edited_docx_file = "static/"+randomName+".docx"  
        edited_docx_path=request.get_host()+"/"+edited_docx_file
        doc.save(edited_docx_file)

     
        return Response({"docx":edited_docx_path})

        

from docx.text.paragraph import Paragraph


def GetParagraphText(paragraph):
    # print("execution")

    def GetTag(element):
        return "%s:%s" % (element.prefix, re.match("{.*}(.*)", element.tag).group(1))
    
    text = ''
    runCount = 0
    documentText=""
    for parent in paragraph._p:
        
        for child in parent:
            tag = GetTag(child)
            # if tag == "w:t":
            #     text += paragraph.runs[runCount].text
            #     runCount += 1
            if tag == "w:hyperlink":
                runCount+=1
                
                for subChild in child:
                    if GetTag(subChild) == "w:r":
                        text+=subChild.text
                        documentText=subChild
                        documentText.text=""

                        # subChild.text=convertLatLngLine(subChild.text)
                
    if runCount>0:
        
        documentText.text=convertLatLngLine(text)
        return getLatLong(text)

    


def getLatLong(line:str):
    expression="^(.+?) - \(Lat:(.+?), Lng:(.+?)\)"
    x=re.search(expression,line)
    horraire= x.group(1)
    lat=float(x.group(2))
    lng=float(x.group(3))
    return (lat,lng)



